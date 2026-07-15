#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_report.py
Genera el informe técnico de 7 páginas del proyecto ETL.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(BASE_DIR, "docs")
os.makedirs(DOCS_DIR, exist_ok=True)

doc = Document()

# Estilos
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# Título principal
title = doc.add_heading("Informe Técnico: Proceso ETL y Dashboard Ejecutivo", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Logística Inteligente del Pacífico S.A.")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].bold = True

meta = doc.add_paragraph("Caso empresarial - Ejercicio Propuesto 2\nJulio 2026")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# 1. Resumen Ejecutivo
doc.add_heading("1. Resumen Ejecutivo", level=1)
doc.add_paragraph(
    "El presente informe describe el diseño, implementación y resultados de un proceso ETL "
    "(Extracción, Transformación y Carga) completamente automatizado para Logística Inteligente "
    "del Pacífico S.A. El objetivo fue consolidar información proveniente de cuatro sistemas "
    "independientes —ERP financiero, sistema GPS, sistema de almacenes y CRM comercial— en un "
    "Data Warehouse centralizado que alimenta un dashboard ejecutivo para la toma de decisiones "
    "estratégicas."
)
doc.add_paragraph(
    "Se identificaron y corrigieron problemas críticos de calidad de datos: registros duplicados, "
    "clientes con múltiples identificadores, formatos inconsistentes de fechas, unidades de medida "
    "heterogéneas y campos incompletos. El resultado es un modelo dimensional robusto con tablas de "
    "hechos financieros, logísticos y de almacén, y una dimensión de clientes."
)

# 2. Análisis de Fuentes de Datos
doc.add_heading("2. Análisis de Fuentes de Datos", level=1)
doc.add_paragraph(
    "La empresa opera con cuatro sistemas independientes que generan datos diarios con diferentes "
    "estructuras y niveles de calidad:"
)

sources = [
    ("ERP Financiero", "Registra facturas, pagos, costos logísticos y vendedores. Campos clave: id_factura, cliente_id, region, fecha_emision, monto_usd, costo_logistico_usd, estado_pago, vendedor_id."),
    ("Sistema GPS", "Controla viajes de vehículos, tiempos de salida/llegada, distancias, combustible y número de entregas. Campos clave: id_viaje, vehiculo_id, fecha, hora_salida, hora_llegada_real, hora_llegada_estimada, distancia_km, combustible_litros."),
    ("Sistema de Almacenes", "Gestiona despachos de productos con cantidades y unidades de medida variadas. Campos clave: id_despacho, cliente_id, producto, cantidad, unidad_medida, fecha_despacho, bodega_origen, peso_bruto."),
    ("CRM Comercial", "Mantiene información de clientes, sectores, regiones comerciales y satisfacción (NPS). Campos clave: id_cliente_crm, nombre_cliente, sector, region_comercial, satisfaccion_nps, estado_cliente."),
]

for name, desc in sources:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{name}: ").bold = True
    p.add_run(desc)

doc.add_paragraph(
    "Cada fuente utiliza convenciones propias para identificadores, fechas y unidades, lo que dificulta "
    "la integración directa y genera diferencias entre reportes financieros y logísticos."
)

# 3. Problemas de Calidad Identificados
doc.add_heading("3. Problemas de Calidad de Datos Identificados", level=1)
problems = [
    "Registros duplicados en ERP (21 registros), GPS (20 registros) y CRM (5 clientes duplicados).",
    "Identificadores de cliente inconsistentes: se encontraron formatos CLI-XXXX, cliente_XXXX y C-XXXX.",
    "Fechas en múltiples formatos: dd/mm/yyyy, yyyy-mm-dd, mm/dd/yyyy, dd-mmm-yyyy y yyyy/mm/dd.",
    "Unidades de medida heterogéneas en almacén: kg, ton, lb, cajas y pallets.",
    "Pesos brutos faltantes en aproximadamente el 5% de los despachos.",
    "Estados de pago faltantes en facturas del ERP.",
    "Encuestas de satisfacción sin respuesta (NPS nulo) en algunos clientes.",
    "Viajes con fechas que excedían el periodo de análisis por problemas de cálculo en origen.",
]
for problem in problems:
    doc.add_paragraph(problem, style="List Bullet")

doc.add_paragraph(
    "Estos problemas fueron corregidos mediante reglas de limpieza, homologación y normalización "
    "programadas en el proceso ETL, garantizando trazabilidad y reproducibilidad."
)

# 4. Arquitectura ETL
doc.add_heading("4. Arquitectura del Proceso ETL", level=1)
doc.add_paragraph(
    "La arquitectura del proceso ETL consta de cuatro capas principales:"
)
architecture = [
    ("Fuentes de datos", "Los cuatro sistemas de origen exportan archivos CSV con la información operativa."),
    ("Capa Staging", "Los archivos CSV se almacenan en data/raw/ sin modificar para conservar el estado original."),
    ("Capa de transformación", "El script etl/etl_process.py ejecuta limpieza, homologación, normalización y cálculo de columnas derivadas."),
    ("Data Warehouse", "Los datos transformados se cargan en SQLite (database/logistica_dw.sqlite) siguiendo un modelo dimensional con tablas de hechos y dimensiones."),
    ("Visualización", "Power BI Desktop consume el Data Warehouse y presenta el dashboard ejecutivo. Adicionalmente se entrega un dashboard en Streamlit."),
]
for title_text, desc in architecture:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{title_text}: ").bold = True
    p.add_run(desc)

doc.add_paragraph(
    "El flujo completo está documentado en diagrams/diagrama_etl.md y puede ejecutarse con un único comando:"
)
doc.add_paragraph("python etl/etl_process.py", style="Intense Quote")

# 5. Proceso de Transformación
doc.add_heading("5. Proceso de Transformación", level=1)
doc.add_paragraph("El proceso de transformación aplica las siguientes reglas por fuente:")

transformations = [
    ("Limpieza de duplicados", "Se eliminan registros completamente idénticos usando drop_duplicates() de pandas."),
    ("Homologación de IDs", "La función homologar_cliente_id() convierte cualquier variante a formato CLI-XXXX mediante expresiones regulares."),
    ("Normalización de fechas", "La función parse_fecha() prueba múltiples formatos hasta obtener un objeto datetime consistente."),
    ("Normalización de unidades", "Todas las cantidades y pesos se convierten a kilogramos: ton×1000, lb×0.453592, cajas×12, pallets×500."),
    ("Imputación de faltantes", "Los pesos faltantes se imputan a partir de la cantidad convertida; los estados de pago faltantes se marcan como Pendiente."),
    ("Columnas calculadas", "Se calculan tiempo de entrega real, diferencia vs. tiempo estimado, indicador de cumplimiento, costo de combustible y categoría NPS."),
]
for title_text, desc in transformations:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{title_text}: ").bold = True
    p.add_run(desc)

# 6. Modelo de Datos y Data Warehouse
doc.add_heading("6. Modelo de Datos y Data Warehouse", level=1)
doc.add_paragraph(
    "El Data Warehouse se implementó en SQLite con un esquema estrella. Las tablas principales son:"
)

tables = [
    ("dim_clientes", "Dimensión", "Catálogo maestro de clientes con sector, región comercial, satisfacción, categoría NPS, ejecutivo de cuenta e ingresos totales."),
    ("fact_financiero", "Hechos", "Transacciones de facturación con cliente, región, fecha, monto, costo logístico y estado de pago."),
    ("fact_logistico", "Hechos", "Viajes con vehículo, fechas, tiempos, distancia, entregas, combustible, cumplimiento y costo de combustible."),
    ("fact_almacen", "Hechos", "Despachos con producto, cantidad original, cantidad normalizada a kg, peso bruto y bodega origen."),
    ("kpis_mensuales", "Agregado", "Indicadores consolidados por mes: ingresos, costos, entregas, viajes, cumplimiento, tiempo promedio."),
    ("kpis_satisfaccion", "Agregado", "Distribución de clientes por categoría NPS."),
]
for table, ttype, desc in tables:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{table} ({ttype}): ").bold = True
    p.add_run(desc)

doc.add_paragraph(
    "Además, se creó la vista v_dashboard para facilitar la conexión desde herramientas de BI. "
    "Los datos procesados también se exportan a CSV y Excel en data/processed/ para uso directo."
)

# 7. Dashboard Ejecutivo y Recomendaciones
doc.add_heading("7. Dashboard Ejecutivo y Recomendaciones", level=1)
doc.add_paragraph(
    "El dashboard ejecutivo incluye los siguientes indicadores y visualizaciones:"
)

kpis = [
    "KPI de ingresos: total y tendencia mensual de ingresos por facturación.",
    "KPI de entregas: total de entregas y número de viajes realizados.",
    "KPI de cumplimiento: porcentaje de viajes entregados a tiempo.",
    "KPI de satisfacción: distribución de NPS (Promotores, Neutros, Detractores).",
    "Costos logísticos: suma de costos logísticos de facturas y combustible.",
    "Tiempo promedio de entrega: evolución mensual en horas.",
    "Segmentación por regiones: ingresos por región comercial.",
    "Tendencias mensuales: análisis de todos los indicadores a lo largo del tiempo.",
]
for kpi in kpis:
    doc.add_paragraph(kpi, style="List Bullet")

doc.add_heading("Recomendaciones para la Alta Dirección", level=2)
recommendations = [
    "Implementar un único master de clientes (MDM) para evitar múltiples identificadores y garantizar la trazabilidad entre sistemas.",
    "Estandarizar los formatos de fecha y unidades de medida en los sistemas de origen para reducir la complejidad del proceso ETL.",
    "Establecer controles de calidad automáticos en los puntos de ingreso de datos (validación de duplicados, campos obligatorios y rangos).",
    "Priorizar la mejora del cumplimiento de entregas, ya que los niveles actuales rondan el 45-50%, lo cual impacta directamente la satisfacción del cliente.",
    "Realizar un análisis de rentabilidad por región y tipo de vehículo para optimizar la asignación de recursos logísticos.",
    "Automatizar la ejecución diaria del ETL mediante un scheduler (por ejemplo, Windows Task Scheduler o Apache Airflow) para mantener el dashboard actualizado.",
    "Adoptar Power BI Service para publicar el dashboard y habilitar el acceso desde dispositivos móviles y la alta dirección.",
    "Definir SLA de respuesta para el área de tecnología ante incidencias en los datos, asegurando la confiabilidad del tablero ejecutivo.",
]
for rec in recommendations:
    doc.add_paragraph(rec, style="List Number")

doc.add_paragraph(
    "\nEl proceso ETL, la base de datos final, el dashboard en Streamlit y esta guía conforman el entregable "
    "completo del ejercicio propuesto. Para el dashboard en Power BI Desktop se proporcionan los datos listos "
    "y las instrucciones de construcción en powerbi/INSTRUCCIONES_POWER_BI.md."
)

# Guardar
output_path = os.path.join(DOCS_DIR, "Informe_Tecnico_ETL.docx")
doc.save(output_path)
print(f"Informe técnico guardado en: {output_path}")
